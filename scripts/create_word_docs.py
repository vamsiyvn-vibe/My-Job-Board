import os
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

PRD_PATH = r"G:\My Drive\.Agents\1_Projects\2026 Job Search\PM_Jobs_Board_PRD.docx"
TDD_PATH = r"G:\My Drive\.Agents\1_Projects\2026 Job Search\PM_Jobs_Board_TDD.docx"

def set_font(run, name="Calibri", size=11, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def format_doc_styles(doc):
    # Set standard margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

def build_prd():
    print(f"Building PRD Word Document at: {PRD_PATH}...")
    doc = docx.Document()
    format_doc_styles(doc)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Product Requirement Document (PRD)\nAutomated PM Job Board")
    set_font(run, name="Calibri", size=22, bold=True)
    
    # Horizontal line replacement
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("*" * 60)
    set_font(run, size=10, bold=False)
    
    # 1. Executive Summary
    h1 = doc.add_heading(level=1)
    run = h1.add_run("1. Executive Summary")
    set_font(run, size=16, bold=True)
    
    p = doc.add_paragraph()
    p.add_run("Finding high-quality Product Management (PM) opportunities at premier companies (Mag 7, AI Labs, High-Growth Startups) is a time-consuming, highly fragmented process. Companies frequently update their internal portals, change job details, or close listings without notifying job aggregators immediately.")
    
    p = doc.add_paragraph()
    p.add_run("This Job Board project solves this by creating a unified, automated, developer-fluent, and self-updating PM Job Leads tracking board. The board runs lightweight, targeted web scraping pipelines against target company career portals, updates a centralized Excel spreadsheet, and renders an interactive, premium web dashboard.")
    
    # 2. Key Problem Statements
    h1 = doc.add_heading(level=1)
    run = h1.add_run("2. Key Problem Statements")
    set_font(run, size=16, bold=True)
    
    problems = [
        ("API and URL Fragmentation: ", "Major platforms (Google, Microsoft, Meta, Amazon, Apple, Nvidia, Netflix) use completely different architectures (Eightfold PCSX, Phenom, Workday, Greenhouse, Ashby, custom GraphQL/batch endpoints), making standard HTML scrapers obsolete quickly."),
        ("Duplicate & Redundant Listings: ", "Traditional job sites often display outdated or duplicate postings. Job hunters need to see only net new listings."),
        ("Ghost Postings: ", "Closed roles continue to sit on job boards long after hiring has ended. Automatic archival is necessary to verify active opportunities."),
        ("Compensation Extraction: ", "Career portals bury salary ranges deep inside text descriptions, requiring candidates to open every listing manually."),
        ("No Level Filtering: ", "Senior, Principal, Group, and Director-level roles are mixed together, wasting search time for experienced PM candidates.")
    ]
    for idx, (title, desc) in enumerate(problems, 1):
        p = doc.add_paragraph(style='List Number')
        r_bold = p.add_run(title)
        set_font(r_bold, bold=True)
        r_desc = p.add_run(desc)
        set_font(r_desc)
        
    # 3. Key Product Features
    h1 = doc.add_heading(level=1)
    run = h1.add_run("3. Key Product Features")
    set_font(run, size=16, bold=True)
    
    features = [
        ("Active Sync Engine: ", "Automatically runs daily scans across Greenhouse, Ashby, Workable, and custom endpoints (Google, Microsoft, Netflix, Nvidia, Apple, Meta, Amazon) to fetch PM listings."),
        ("Unified Status Workflow: ", "Allows labeling jobs as Lead, Consideration (Double-click lists to view details), or Archived."),
        ("Archival Tracking & Reason-Mapping: ", "Automatically archives roles that disappear from career portals, classifying them as Closed, while user-triggered archives are marked as User Archived."),
        ("Smart Compensation Extractor: ", "Employs a multi-format regex parser to extract salary ranges from job descriptions, displaying them as columns in lists and highlights in cards."),
        ("Cohort-Based Level Filtering: ", "Standardizes filters by target company cohorts (e.g. Mag 7 vs. Non-Mag 7 vs. High-Growth Startups vs. AI Labs) and supports strict level filtering (e.g. restricting Amazon to Principal PMs and above, and supporting a 'Generic' target level for entries with no seniority titles)."),
        ("Dynamic Company Feeds: ", "Enables users to add new companies dynamically with auto-platform and board ID detection, immediately triggering a single-company scraper run and automatic data enrichment.")
    ]
    for title, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        r_bold = p.add_run(title)
        set_font(r_bold, bold=True)
        r_desc = p.add_run(desc)
        set_font(r_desc)
        
    # 4. System Design & User Experience
    h1 = doc.add_heading(level=1)
    run = h1.add_run("4. System Design & User Experience")
    set_font(run, size=16, bold=True)
    
    doc.add_paragraph("The application is designed around a premium developer-style dashboard. Users can toggle between Grid Card layouts and dense spreadsheet List layouts, filter job listings by company name, location, cohort role level, and salary range (e.g. min $150k+, min $200k+, etc.). Double-clicking on a job card opens a slide-over modal containing the live scraped job description and a custom PM Interview Loop preparation guide customized for that company.")
    
    # 5. System Limitations & Mitigations
    h1 = doc.add_heading(level=1)
    run = h1.add_run("5. System Limitations & Mitigations")
    set_font(run, size=16, bold=True)
    
    limits = [
        ("Anti-Scraping / Cloudflare: ", "Sites like Meta block direct requests. Mitigation: We use session-based cookie persistence and browser-impersonation headers (Sec-Ch-Ua, Sec-Fetch-Mode) to avoid 400 Bad Request errors."),
        ("Rate Limits: ", "Rapid queries can cause IP bans. Mitigation: The scrapers implement sleep delays (e.g., 0.2s - 0.5s) between consecutive detail page queries."),
        ("Schema Drift: ", "If a company updates its GraphQL query ID or changes its endpoint, the scraper will fail. Mitigation: The architecture falls back to generic DuckDuckGo search indexes if the main API endpoint throws an exception.")
    ]
    for title, desc in limits:
        p = doc.add_paragraph(style='List Bullet')
        r_bold = p.add_run(title)
        set_font(r_bold, bold=True)
        r_desc = p.add_run(desc)
        set_font(r_desc)
        
    doc.save(PRD_PATH)
    print("PRD Word Document saved successfully.")

def build_tdd():
    print(f"Building TDD Word Document at: {TDD_PATH}...")
    doc = docx.Document()
    format_doc_styles(doc)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Technical Design Document (TDD)\nAutomated PM Job Board")
    set_font(run, name="Calibri", size=22, bold=True)
    
    # Horizontal line replacement
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("*" * 60)
    set_font(run, size=10, bold=False)
    
    # 1. System Architecture Overview
    h1 = doc.add_heading(level=1)
    run = h1.add_run("1. System Architecture Overview")
    set_font(run, size=16, bold=True)
    
    doc.add_paragraph("The application consists of three main components: a crawler/scraper script, a central Excel-based database, and a local Flask API server that serves the frontend UI. The application is completely local, ensuring all credentials, spreadsheets, and scraper configurations are stored securely on the user's G: Drive.")
    
    # Add Table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    headers = ["Component", "File / Tech Stack", "Responsibilities"]
    for i, header in enumerate(headers):
        hdr_cell = table.rows[0].cells[i]
        hdr_cell.text = header
        set_font(hdr_cell.paragraphs[0].runs[0], bold=True)
        
    components_data = [
        ("Database", "Job_Leads_Tracker.xlsx", "Stores all job listings, cohorts, URLs, salaries, notes, and application status. Serves as the single source of truth."),
        ("Scraper Engine", "scripts/check_new_jobs.py", "Executes web crawlers for all listed companies, filters listings, extracts salary ranges, and appends rows to Excel."),
        ("Flask API Server", "board/server.py", "Serves UI static assets, reads Excel database with caching, handles status updates, scrapes job descriptions on-the-fly, and triggers background crawls."),
        ("Web Dashboard", "HTML5 / CSS3 / Vanilla JS (ES6+)", "Provides responsive dark-themed UI. Features card and list layouts, search filtering, and slide-over detailed job description modals with interview guides.")
    ]
    for idx, row_data in enumerate(components_data, 1):
        for col_idx, text in enumerate(row_data):
            cell = table.rows[idx].cells[col_idx]
            cell.text = text
            if col_idx == 0:
                set_font(cell.paragraphs[0].runs[0], bold=True)
            else:
                set_font(cell.paragraphs[0].runs[0])
                
    doc.add_paragraph() # space after table
    
    # 2. Database Schema (Excel Rows Mapping)
    h1 = doc.add_heading(level=1)
    run = h1.add_run("2. Database Schema (Excel Rows Mapping)")
    set_font(run, size=16, bold=True)
    
    doc.add_paragraph("The Excel sheet 'Job Leads' uses a 15-column schema, structured as follows:")
    
    schema_table = doc.add_table(rows=16, cols=4)
    schema_table.style = 'Table Grid'
    
    s_headers = ["Col #", "Column Name", "Data Type", "Description / Sample Value"]
    for i, sh in enumerate(s_headers):
        hdr_cell = schema_table.rows[0].cells[i]
        hdr_cell.text = sh
        set_font(hdr_cell.paragraphs[0].runs[0], bold=True)
        
    schema_data = [
        ("1", "Select to Apply", "String (Checkbox)", "[ ] or [x] (mapped to Consideration status)"),
        ("2", "Cohort", "String", "Mag 7, Non-Mag 7, High-Growth Startups, AI Labs"),
        ("3", "Company", "String", "Company name (e.g. Apple, Greenhouse Company)"),
        ("4", "Role Title", "String", "Job title (e.g. Senior Product Manager)"),
        ("5", "Location", "String", "Location string or Remote"),
        ("6", "Key Focus", "String", "Extracted domain focus (e.g. Growth, Core Tech)"),
        ("7", "Status", "String", "Lead, Consideration, Archived"),
        ("8", "URL", "String (Hyperlink)", "The application portal link"),
        ("9", "Date Added", "String (Date)", "Format: YYYY-MM-DD"),
        ("10", "Notes", "String", "Additional custom comments"),
        ("11", "Compensation", "String", "Normalized range (e.g. $185k-$240k or N/A)"),
        ("12", "App Status", "String", "Application state (e.g. Not Applied, Applied, Interviewing)"),
        ("13", "App Outcome", "String", "Application result (e.g. Active / Pending, Rejected)"),
        ("14", "Resume Link", "String", "File path/link to the resume submitted"),
        ("15", "Archive Reason", "String", "Reason (e.g. Closed, User Archived)")
    ]
    for idx, row_data in enumerate(schema_data, 1):
        for col_idx, text in enumerate(row_data):
            cell = schema_table.rows[idx].cells[col_idx]
            cell.text = text
            if col_idx in [0, 1]:
                set_font(cell.paragraphs[0].runs[0], bold=True)
            else:
                set_font(cell.paragraphs[0].runs[0])
                
    doc.add_paragraph() # space after table
    
    # 3. Scraper Specifications (Per Company)
    h1 = doc.add_heading(level=1)
    run = h1.add_run("3. Scraper Specifications (Per Company)")
    set_font(run, size=16, bold=True)
    
    scrapers = [
        ("Google: ", "Simulated HTTP POST payload hitting Google's batchexecute endpoint. Parses internal job lists and extracts titles and locations."),
        ("Microsoft: ", "Queries the Microsoft Careers Eightfold search API (GET /api/pcsx/search) after establishing a session and resolving CSRF tokens."),
        ("Nvidia: ", "Queries Nvidia's Eightfold search API similar to Microsoft, applying correct headers and session states."),
        ("Netflix: ", "Crawls the Phenom-powered portal. Parses the HTML response and extracts pre-rendered job lists from the <code id=\"smartApplyData\"> script block."),
        ("Apple: ", "Accesses Apple careers website. Parses and hydrates variables injected in script blocks directly in Apple's HTML response."),
        ("Meta: ", "Direct GraphQL call to /api/graphql/ with doc ID 26703205452636175, utilizing session cookies and browser headers to prevent blocks."),
        ("Amazon: ", "Queries Amazon's public careers search API (GET /en/search.json) with dynamically built search terms based on user's target levels."),
        ("Standard ATS (Greenhouse, Ashby, Workable): ", "Auto-detected from URL parameters; parses job list JSON/HTML directly.")
    ]
    for title, desc in scrapers:
        p = doc.add_paragraph(style='List Bullet')
        r_bold = p.add_run(title)
        set_font(r_bold, bold=True)
        r_desc = p.add_run(desc)
        set_font(r_desc)
        
    # 4. Core Engineering Flows
    h1 = doc.add_heading(level=1)
    run = h1.add_run("4. Core Engineering Flows")
    set_font(run, size=16, bold=True)
    
    flows = [
        ("4.1 Performance Caching (JOBS_CACHE): ", "To avoid reading the Excel spreadsheet on every single user interaction or filter update, server.py loads jobs into a global in-memory cache variable JOBS_CACHE. On every API query, it checks the spreadsheet modification time (mtime) and only reloads the Excel file if the file on disk has changed."),
        ("4.2 Immediate Scan & Company Enrichment: ", "When a new company is added in the UI, the backend creates a database entry in config.json, parses the portal type, and spawns a background thread running the scraper script specifically for that company (e.g., check_new_jobs.py \"Company Name\"). Simultaneously, it queries DuckDuckGo search indexes to automatically discover company metadata (HQ location, foundation year, revenue, employee count, and company domain) to populate the dashboard."),
        ("4.3 Smart Compensation Parser: ", "When scraping a job description, a regular expression parser runs to extract salary numbers. It processes expressions like $150,000 - $220,000, $180k to $250k, or $120/hr, and normalizes them into standard $Xk-$Yk formats. This details is written to Column 11 of the spreadsheet and rendered directly on the dashboard cards."),
        ("4.4 Cohort-Based Level Filtering: ", "Job titles are filtered against the target company's cohort settings. When a crawl executes, title_matches_levels checks if the job title contains level-specific whitelisted phrases or falls back to 'Generic' if the title contains no seniority suffixes (like Senior, Lead, Principal, Staff, Group, Director, VP, Head, II, etc.)."),
        ("4.5 Deletion Cascade: ", "Deleting a monitored company from the feeds widget triggers a DELETE API request. The backend removes the company config from config.json, iterates backwards through rows in the Excel spreadsheet to delete all job rows belonging to that company, and clears the in-memory cache to sync the UI immediately.")
    ]
    for title, desc in flows:
        p = doc.add_paragraph()
        r_bold = p.add_run(title)
        set_font(r_bold, bold=True)
        r_desc = p.add_run(desc)
        set_font(r_desc)
        
    doc.save(TDD_PATH)
    print("TDD Word Document saved successfully.")

if __name__ == "__main__":
    build_prd()
    build_tdd()
    print("Success! Created word documents.")
